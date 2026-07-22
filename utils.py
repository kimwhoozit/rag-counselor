import os
from typing import List
import pandas as pd
import docx
from pypdf import PdfReader
import olefile

def extract_text_via_gemini(file_path: str, api_key) -> str:
    """Uploads a file to Gemini API to extract its text content using multimodal capabilities."""
    import google.generativeai as genai
    from llm_service import get_working_keys
    
    keys = get_working_keys(api_key)
    if not keys:
        raise ValueError("Gemini API Key가 설정되어 있지 않아 OCR Fallback을 수행할 수 없습니다.")
    
    # Configure with the first working key
    genai.configure(api_key=keys[0])
    
    try:
        uploaded_file = genai.upload_file(path=file_path)
    except Exception as upload_err:
        raise RuntimeError(f"Gemini API 파일 임시 업로드 실패: {str(upload_err)}")
        
    try:
        # Prompt model to extract all text
        # gemini-1.5-flash is stable and supports PDF/Image/Office document uploads
        model = genai.GenerativeModel("gemini-1.5-flash")
        prompt = (
            "이 문서 파일의 모든 텍스트 내용을 한 글자도 빠짐없이 있는 그대로 추출해 주세요.\n"
            "요약하거나 설명하지 말고, 오직 문서 안의 텍스트 본문만 출력해야 합니다."
        )
        
        response = model.generate_content([uploaded_file, prompt])
        text = response.text
        if not text or not text.strip():
            raise ValueError("Gemini OCR에서 반환된 텍스트가 비어 있습니다.")
        return text
    finally:
        # Clean up the file from Gemini storage immediately
        try:
            genai.delete_file(uploaded_file.name)
        except Exception:
            pass

def parse_file(file_path: str, api_key = None) -> str:
    """Reads various file types and returns their text contents.
    Supports .txt, .md, .docx, .xlsx, .xls, and .pdf files.
    """
    import os
    from typing import Any
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext in ['.txt', '.md']:
            # Try multiple encodings for Korean Windows environments
            encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-16', 'latin-1']
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode the text file {file_path} with standard encodings.")
            
        elif ext == '.docx':
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Include table contents
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            parsed = '\n'.join(full_text)
            
            if (not parsed.strip() or len(parsed.strip()) < 100) and api_key:
                try:
                    return extract_text_via_gemini(file_path, api_key)
                except Exception:
                    pass
            return parsed
            
        elif ext in ['.xlsx', '.xls']:
            try:
                excel_data = pd.read_excel(file_path, sheet_name=None)
                full_text = []
                for sheet_name, df in excel_data.items():
                    # Drop rows that are completely empty to save space
                    df_cleaned = df.dropna(how='all')
                    if not df_cleaned.empty:
                        full_text.append(f"--- Sheet: {sheet_name} ---")
                        # Convert to CSV format: much more compact than df.to_string() 
                        # and extremely easy for LLMs to understand without padding overhead.
                        full_text.append(df_cleaned.to_csv(index=False))
                return '\n'.join(full_text)
            except Exception as e:
                raise ValueError(f"Failed to read Excel file: {str(e)}")
                
        elif ext == '.pdf':
            reader = PdfReader(file_path)
            full_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
            parsed = '\n'.join(full_text)
            
            if (not parsed.strip() or len(parsed.strip()) < 100) and api_key:
                try:
                    return extract_text_via_gemini(file_path, api_key)
                except Exception as ocr_err:
                    if parsed.strip():
                        return parsed
                    raise ValueError(f"PDF에서 디지털 텍스트 추출 결과가 비어 있어 Gemini OCR을 시도했으나 실패했습니다: {str(ocr_err)}")
            return parsed
            
        elif ext == '.hwpx':
            parsed = extract_hwpx_text(file_path)
            if (not parsed.strip() or len(parsed.strip()) < 100) and api_key:
                try:
                    return extract_text_via_gemini(file_path, api_key)
                except Exception:
                    pass
            return parsed
            
        elif ext == '.hwp':
            parsed = extract_hwp_text(file_path)
            if (not parsed.strip() or len(parsed.strip()) < 100) and api_key:
                try:
                    return extract_text_via_gemini(file_path, api_key)
                except Exception:
                    pass
            return parsed
            
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
    except Exception as main_err:
        if api_key and ext in ['.pdf', '.docx', '.xlsx', '.xls']:
            try:
                return extract_text_via_gemini(file_path, api_key)
            except Exception as ocr_err:
                raise ValueError(f"문서 기본 파싱 실패 및 Gemini OCR 실패: {str(main_err)} / OCR 오류: {str(ocr_err)}")
        raise main_err

def extract_hwpx_text(file_path: str) -> str:
    """Extracts text from HWPX (Hancom Office HWPX) files by unzipping 
    and reading content section XMLs.
    """
    import zipfile
    import xml.etree.ElementTree as ET
    
    text_content = []
    with zipfile.ZipFile(file_path) as z:
        # Find all section xml files under Contents/ and sort them
        sections = sorted([
            name for name in z.namelist()
            if name.startswith("Contents/section") and name.endswith(".xml")
        ])
        
        for section in sections:
            xml_data = z.read(section)
            root = ET.fromstring(xml_data)
            
            section_texts = []
            for elem in root.iter():
                # Elements in HWPX XML usually have namespaces, e.g. {namespace}t
                tag_name = elem.tag.split('}')[-1]
                if tag_name == 't' and elem.text:
                    section_texts.append(elem.text)
            
            # If for some reason HWPX doesn't use the standard HWPML structure,
            # fallback to gathering all readable text elements
            if not section_texts:
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        txt = elem.text.strip()
                        # simple filter to ignore layout numbers and short characters
                        if len(txt) > 1:
                            section_texts.append(txt)
                            
            if section_texts:
                text_content.append('\n'.join(section_texts))
                
    return '\n\n'.join(text_content)

def extract_hwp_text(file_path: str) -> str:
    """Extracts text from older HWP (Hancom Office 5.0+) files by reading
    the PrvText (preview text) stream using the olefile package.
    """
    try:
        f = olefile.OleFileIO(file_path)
        if f.exists('PrvText'):
            data = f.openstream('PrvText').read()
            # HWP5 PrvText stream is typically stored as UTF-16 LE
            text = data.decode('utf-16-le', errors='ignore')
            return text
        else:
            raise ValueError("PrvText stream not found inside the HWP file structure.")
    except Exception as e:
        raise ValueError(f"OLE extraction failed: {str(e)}")


def split_text(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[str]:
    """Splits a long text into chunks of chunk_size characters with chunk_overlap.
    Splits are made on paragraph boundaries (new lines) if possible to preserve context.
    """
    if not text:
        return []
        
    # Split text into paragraphs/lines
    lines = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for line in lines:
        line_len = len(line) + 1  # count the newline char
        if current_length + line_len > chunk_size:
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
            
            # Simple overlap logic: keep last few lines that sum up to less than chunk_overlap
            overlap_chunk = []
            overlap_len = 0
            for l in reversed(current_chunk):
                if overlap_len + len(l) + 1 < chunk_overlap:
                    overlap_chunk.insert(0, l)
                    overlap_len += len(l) + 1
                else:
                    break
            
            current_chunk = overlap_chunk
            current_length = overlap_len
            
        current_chunk.append(line)
        current_length += line_len
        
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
        
    # If any single chunk is still too large (e.g. no newlines), split it character-wise
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > chunk_size * 1.5:
            start = 0
            while start < len(chunk):
                end = min(start + chunk_size, len(chunk))
                final_chunks.append(chunk[start:end])
                start += chunk_size - chunk_overlap
        else:
            final_chunks.append(chunk)
            
    return [c.strip() for c in final_chunks if c.strip()]
